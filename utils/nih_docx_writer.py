# utils/nih_docx_writer.py
from __future__ import annotations

import re
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Any

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


# DOCX anchors (instruction paragraphs in the NIH template)
# NOTE: Kept same structure, but added whitespace/linebreak-safe matching.
PROMPT_ANCHORS: List[Tuple[str, List[str]]] = [
    ("e1_q1", ["Summarize the types and estimated amount of scientific data expected to be generated in the project"]),
    ("e1_q2", ["Describe which scientific data from the project will be preserved and shared and provide the rationale for this decision"]),
    ("e1_q3", ["Briefly list the metadata, other relevant data, and any associated documentation"]),
    ("e2_q1", ["State whether specialized tools, software, and/or code are needed to access or manipulate shared scientific data"]),
    ("e3_q1", ["State what common data standards will be applied to the scientific data and associated metadata to enable interoperability of datasets and resources"]),
    ("e4_q1", ["Provide the name of the repository(ies) where scientific data and metadata arising from the project will be archived"]),
    ("e4_q2", ["Describe how the scientific data will be findable and identifiable"]),
    ("e4_q3", ["Describe when the scientific data will be made available to other users"]),
    ("e5_q1", ["NIH expects that in drafting Plans, researchers maximize the appropriate sharing of scientific data."]),
    ("e5_q2", ["State whether access to the scientific data will be controlled"]),
    ("e5_q3", ["If generating scientific data derived from humans, describe how the privacy, rights, and confidentiality of human research participants will be protected"]),
    ("e6_q1", ["Describe how compliance with this Plan will be monitored and managed, frequency of oversight, and by whom at your institution"]),
]


# ---------- DOCX helpers ----------

def _insert_paragraph_after(paragraph: Paragraph, text: str = "", style_name: Optional[str] = None) -> Paragraph:
    """Insert a new paragraph after an existing paragraph (python-docx safe)."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)

    if style_name:
        try:
            new_para.style = style_name
        except Exception:
            pass

    if text:
        new_para.add_run(text)

    return new_para


def _delete_paragraph(paragraph: Paragraph) -> None:
    """Remove a paragraph from the document."""
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None  # type: ignore


def _normalize_ws_after(doc: Document, idx: int) -> None:
    """Delete ALL consecutive blank paragraphs starting at idx."""
    j = idx
    while j < len(doc.paragraphs):
        p = doc.paragraphs[j]
        if (p.text or "").strip() == "":
            _delete_paragraph(p)
            continue
        break


def _norm_ws(s: str) -> str:
    """
    Normalize whitespace to make anchor matching robust against:
      - tabs
      - newlines
      - multiple spaces
      - Word wrapping
    """
    return re.sub(r"\s+", " ", (s or "").strip()).lower()


def _find_anchor_paragraph_index(doc: Document, anchor_aliases: List[str]) -> Optional[int]:
    """
    Robust anchor finding:
      - compares normalized (collapsed) whitespace
      - case-insensitive
    """
    aliases_norm = [_norm_ws(a) for a in (anchor_aliases or []) if (a or "").strip()]
    if not aliases_norm:
        return None

    for i, p in enumerate(doc.paragraphs):
        pt = _norm_ws(p.text or "")
        if not pt:
            continue
        for a in aliases_norm:
            if a and a in pt:
                return i
    return None


def _remove_placeholder_only_paragraphs(doc: Document) -> None:
    placeholders = {k.lower() for k, _ in PROMPT_ANCHORS}
    for p in doc.paragraphs:
        if (p.text or "").strip().lower() in placeholders:
            p.text = ""


def _copy_paragraph_format(dst: Paragraph, src: Paragraph) -> None:
    """
    Copy indent/spacing/alignment from src paragraph to dst paragraph so the answer aligns
    exactly like the NIH template.
    """
    dpf = dst.paragraph_format
    spf = src.paragraph_format

    # Indentation
    dpf.left_indent = spf.left_indent
    dpf.right_indent = spf.right_indent
    dpf.first_line_indent = spf.first_line_indent

    # Spacing
    dpf.space_before = spf.space_before
    dpf.space_after = spf.space_after
    dpf.line_spacing = spf.line_spacing

    # Alignment / pagination
    dpf.alignment = spf.alignment
    dpf.keep_together = spf.keep_together
    dpf.keep_with_next = spf.keep_with_next
    dpf.page_break_before = spf.page_break_before
    dpf.widow_control = spf.widow_control


def _add_answer_paragraph_after(
    anchor: Paragraph,
    text: str,
    style_name: Optional[str],
    format_source: Paragraph,
) -> Paragraph:
    """
    Add a paragraph after anchor, using:
      - style_name (usually the template answer line style)
      - paragraph_format copied from format_source (indentation/spacing)
    """
    p = _insert_paragraph_after(anchor, "", style_name=style_name)
    _copy_paragraph_format(p, format_source)
    if text:
        p.add_run(text)
    return p


def _write_block_after_anchor(doc: Document, anchor_aliases: List[str], answer: str) -> None:
    """
    Insert answer after the NIH prompt, formatted like the template:
      - normal paragraphs aligned under the prompt (indent copied)
      - exactly one blank line between paragraphs and between sections
      - list lines ("- ...") become normal paragraphs (no bullet indentation)
    """
    idx = _find_anchor_paragraph_index(doc, anchor_aliases)
    if idx is None:
        return

    anchor_p = doc.paragraphs[idx]

    # Prefer style/format from the template's answer line (often the blank paragraph right after the prompt)
    style_name = anchor_p.style.name if anchor_p.style else None
    format_source = anchor_p

    if idx + 1 < len(doc.paragraphs):
        nxt = doc.paragraphs[idx + 1]
        # If the template has a blank "answer line" paragraph, copy its style/format
        if (nxt.text or "").strip() == "":
            format_source = nxt
            if nxt.style:
                style_name = nxt.style.name

    # Remove ALL blank paragraphs right after the anchor prompt
    _normalize_ws_after(doc, idx + 1)

    answer = (answer or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if not answer:
        # keep one blank answer line (matches template feel)
        _add_answer_paragraph_after(anchor_p, "", style_name, format_source)
        return

    # Turn the answer into a list of paragraphs separated by blank lines.
    # Also turn "- item" / "• item" into its own paragraph "item" (no bullet style).
    raw_lines = [ln.rstrip() for ln in answer.split("\n")]
    paragraphs: List[str] = []
    buf: List[str] = []

    def flush_buf() -> None:
        nonlocal buf
        joined = " ".join([x.strip() for x in buf if x.strip()]).strip()
        if joined:
            paragraphs.append(joined)
        buf = []

    for ln in raw_lines:
        s = ln.strip()
        if s == "":
            flush_buf()
            continue

        if s.startswith("- "):
            flush_buf()
            paragraphs.append(s[2:].strip())
            continue

        if s.startswith("• "):
            flush_buf()
            paragraphs.append(s[2:].strip())
            continue

        buf.append(s)

    flush_buf()

    # Write paragraphs with EXACTLY one blank line between them
    last_p = anchor_p
    for i, para_text in enumerate(paragraphs):
        last_p = _add_answer_paragraph_after(last_p, para_text, style_name, format_source)
        if i != len(paragraphs) - 1:
            last_p = _add_answer_paragraph_after(last_p, "", style_name, format_source)

    # one blank line after the whole block (between sections)
    _add_answer_paragraph_after(last_p, "", style_name, format_source)


# ---------- Markdown parsing (legacy fallback) ----------

def _strip_markdown_keep_structure(md: str) -> str:
    if not md:
        return ""
    text = md.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)  # [text](url) -> text
    text = text.replace("**", "").replace("*", "")        # remove emphasis markers
    text = re.sub(r"\n{3,}", "\n\n", text)                # collapse excessive blank lines
    return text.strip()


def _split_into_heading_blocks(text: str) -> List[Tuple[str, str]]:
    lines = text.split("\n")
    blocks: List[Tuple[str, List[str]]] = []

    def is_heading(line: str) -> bool:
        s = line.strip()
        if not s:
            return False
        if s.lower().startswith("element "):
            return True
        if s.startswith("### "):
            return True
        return False

    current_heading: Optional[str] = None
    current_content: List[str] = []

    for line in lines:
        if is_heading(line):
            if current_heading is not None:
                blocks.append((current_heading, current_content))
            current_heading = line.strip().lstrip("#").strip()
            current_content = []
        else:
            current_content.append(line)

    if current_heading is not None:
        blocks.append((current_heading, current_content))

    out: List[Tuple[str, str]] = []
    for h, content_lines in blocks:
        content = "\n".join(content_lines).strip()
        out.append((h, content))
    return out


def _extract_blocks_from_generated_markdown(generated_markdown: str) -> Dict[str, str]:
    plain = _strip_markdown_keep_structure(generated_markdown)
    heading_blocks = _split_into_heading_blocks(plain)

    hmap: Dict[str, str] = {}
    for h, c in heading_blocks:
        hmap[h.lower()] = c.strip()

    def get_h(contains: str) -> str:
        contains = contains.lower()
        for h in hmap.keys():
            if contains in h:
                return hmap[h]
        return ""

    blocks: Dict[str, str] = {}

    # Element 1 sub-headings
    blocks["e1_q1"] = get_h("types and amount of scientific data expected")
    blocks["e1_q2"] = get_h("scientific data that will be preserved and shared")
    blocks["e1_q3"] = get_h("metadata, other relevant data")

    # Element 2/3/6 whole blocks
    blocks["e2_q1"] = get_h("element 2: related tools")
    blocks["e3_q1"] = get_h("element 3: standards")
    blocks["e6_q1"] = get_h("element 6: oversight")

    # Element 4 sub-headings
    blocks["e4_q1"] = get_h("repository where scientific data and metadata will be archived")
    blocks["e4_q2"] = get_h("how scientific data will be findable and identifiable")
    blocks["e4_q3"] = get_h("when and how long the scientific data will be made available")

    # Element 5 sub-headings
    blocks["e5_q1"] = get_h("factors affecting subsequent access")
    blocks["e5_q2"] = get_h("whether access to scientific data will be controlled")
    blocks["e5_q3"] = get_h("protections for privacy, rights, and confidentiality")

    for k, v in list(blocks.items()):
        blocks[k] = (v or "").strip().lstrip(" \t\n\r:.-").strip()

    return blocks


# ---------- Plan JSON helpers (NEW) ----------

def _extract_blocks_from_plan_json(plan_json: Dict[str, Any]) -> Dict[str, str]:
    """
    Canonical mapping: plan_json["answers"] keys -> docx anchor keys.
    Expected answer keys:
      1.1, 1.2, 1.3, 2.1, 3.1, 4.1, 4.2, 4.3, 5.1, 5.2, 5.3, 6.1
    """
    answers = (plan_json or {}).get("answers") or {}

    def g(k: str) -> str:
        v = answers.get(k)
        return (v or "").strip()

    return {
        "e1_q1": g("1.1"),
        "e1_q2": g("1.2"),
        "e1_q3": g("1.3"),
        "e2_q1": g("2.1"),
        "e3_q1": g("3.1"),
        "e4_q1": g("4.1"),
        "e4_q2": g("4.2"),
        "e4_q3": g("4.3"),
        "e5_q1": g("5.1"),
        "e5_q2": g("5.2"),
        "e5_q3": g("5.3"),
        "e6_q1": g("6.1"),
    }


# ---------- Public API ----------

def build_nih_docx_from_template(
    template_docx_path: str,
    output_docx_path: str,
    project_title: str,  # kept for backward compatibility; not inserted anymore
    generated_markdown: str = "",
    *,
    plan_json: Optional[Dict[str, Any]] = None,
) -> None:
    """
    Preferred path:
      - if plan_json is provided, use plan_json["answers"] (no markdown parsing)
    Fallback path:
      - parse generated_markdown (legacy)
    """
    doc = Document(str(template_docx_path))

    # DO NOT insert "Project Title: ..." under the heading anymore.

    if plan_json:
        blocks = _extract_blocks_from_plan_json(plan_json)
    else:
        blocks = _extract_blocks_from_generated_markdown(generated_markdown or "")

    _remove_placeholder_only_paragraphs(doc)

    for key, aliases in PROMPT_ANCHORS:
        answer = (blocks.get(key, "") or "").strip()
        if not answer:
            continue
        _write_block_after_anchor(doc, aliases, answer)

    Path(str(output_docx_path)).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx_path))
